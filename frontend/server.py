import asyncio
import json
import logging
import mimetypes
import shutil
import time as _time
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

import yaml
from fastapi import FastAPI, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles

app = FastAPI(title="ViMax Preview")
logger = logging.getLogger("vimax.server")

BASE_DIR = Path(__file__).parent.parent
WORKING_DIR = BASE_DIR / ".working_dir"
STATIC_DIR = Path(__file__).parent / "static"
CONFIGS_DIR = BASE_DIR / "configs"

SCENE_REF_EXTS = {".png", ".jpg", ".jpeg", ".webp"}

app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


# ── Config 发现 ──────────────────────────────────────────────────────────────

_config_cache: Dict[str, str] = {}  # project_name -> config_path


def find_config(project: str) -> Optional[str]:
    if project in _config_cache:
        return _config_cache[project]
    for yaml_file in CONFIGS_DIR.glob("*.yaml"):
        try:
            cfg = yaml.safe_load(yaml_file.read_text(encoding="utf-8"))
            if Path(cfg.get("working_dir", "")).name == project:
                _config_cache[project] = str(yaml_file)
                return str(yaml_file)
        except Exception:
            pass
    return None


# ── Pipeline 缓存 ────────────────────────────────────────────────────────────

_pipeline_cache: Dict[str, tuple] = {}  # project_name -> (fingerprint, pipeline instance)


def _apply_api_overrides(config: dict, overrides: dict):
    """Apply api_overrides from metadata onto a base config dict (in-place)."""
    if "chat_model" in overrides:
        for k, v in overrides["chat_model"].items():
            if v is not None and v != "":
                config.setdefault("chat_model", {}).setdefault("init_args", {})[k] = v
    if "image_generator" in overrides:
        ov = overrides["image_generator"]
        if ov.get("class_path"):
            config["image_generator"]["class_path"] = ov["class_path"]
        if ov.get("init_args"):
            merged = dict(config.get("image_generator", {}).get("init_args", {}))
            for k, v in ov["init_args"].items():
                if v is not None and v != "":
                    merged[k] = v
            config.setdefault("image_generator", {})["init_args"] = merged
    if "video_generator" in overrides:
        ov = overrides["video_generator"]
        if ov.get("class_path"):
            config["video_generator"]["class_path"] = ov["class_path"]
        if ov.get("init_args"):
            merged = dict(config.get("video_generator", {}).get("init_args", {}))
            for k, v in ov["init_args"].items():
                if v is not None and v != "":
                    merged[k] = v
            config.setdefault("video_generator", {})["init_args"] = merged


def get_pipeline(project: str):
    wd = WORKING_DIR / project
    meta_path = wd / "metadata.json"
    overrides: dict = {}
    if meta_path.exists():
        try:
            overrides = json.loads(meta_path.read_text(encoding="utf-8")).get("api_overrides", {})
        except Exception:
            pass
    fingerprint = json.dumps(overrides, sort_keys=True)

    cached = _pipeline_cache.get(project)
    if cached and cached[0] == fingerprint:
        return cached[1]

    config_path = find_config(project)
    if not config_path:
        raise HTTPException(404, f"No config file found for project '{project}'")

    with open(config_path, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    _apply_api_overrides(config, overrides)

    from langchain.chat_models import init_chat_model
    from tools.render_backend import RenderBackend
    from pipelines.script2video_pipeline import Script2VideoPipeline

    chat_model = init_chat_model(**config["chat_model"]["init_args"])
    backend = RenderBackend.from_config(config)
    pipeline = Script2VideoPipeline(
        chat_model=chat_model,
        image_generator=backend.image_generator,
        video_generator=backend.video_generator,
        working_dir=config["working_dir"],
    )
    _pipeline_cache[project] = (fingerprint, pipeline)
    logger.info(f"Pipeline (re)loaded for project '{project}' using config '{config_path}'")
    return pipeline


# ── Task 系统 ────────────────────────────────────────────────────────────────

_tasks: Dict[str, Dict] = {}  # task_id -> {status, error_msg}


def _new_task() -> str:
    tid = str(uuid.uuid4())
    _tasks[tid] = {"status": "running", "error_msg": None}
    return tid


def _task_done(tid: str):
    _tasks[tid]["status"] = "done"
    _tasks[tid]["status_msg"] = ""


def _task_error(tid: str, msg: str):
    _tasks[tid]["status"] = "error"
    _tasks[tid]["error_msg"] = msg


def _task_progress(tid: str, msg: str):
    _tasks[tid]["status_msg"] = msg


# ── Frame events 工具 ────────────────────────────────────────────────────────

def _setup_frame_events(pipeline, wd: Path, target_shot_idx: int, target_frame_type: str):
    """为 generate_frame_for_single_shot 预初始化所有 frame_events。
    目标帧的 event 保持 unset；其他已存在的帧 event 置为 set。"""
    from pipelines.script2video_pipeline import Script2VideoPipeline
    Script2VideoPipeline.frame_events = {}  # 清空类级别状态
    shots_dir = wd / "shots"
    for d in sorted(shots_dir.iterdir()):
        if not d.is_dir():
            continue
        idx = int(d.name)
        events = {}
        for ft in ("first_frame", "last_frame"):
            ev = asyncio.Event()
            # 目标帧保持 unset（让方法正常写入后 set）
            if not (idx == target_shot_idx and ft == target_frame_type):
                if (d / f"{ft}.png").exists():
                    ev.set()
            events[ft] = ev
        Script2VideoPipeline.frame_events[idx] = events


def _preset_all_frame_events(pipeline, wd: Path):
    """为 generate_video_for_single_shot 预初始化：所有已有帧的 event 都置为 set。"""
    from pipelines.script2video_pipeline import Script2VideoPipeline
    Script2VideoPipeline.frame_events = {}
    shots_dir = wd / "shots"
    for d in sorted(shots_dir.iterdir()):
        if not d.is_dir():
            continue
        idx = int(d.name)
        events = {}
        for ft in ("first_frame", "last_frame"):
            ev = asyncio.Event()
            if (d / f"{ft}.png").exists():
                ev.set()
            events[ft] = ev
        Script2VideoPipeline.frame_events[idx] = events


def _get_first_shot_ff_pair(wd: Path, shot_idx: int):
    """找到该 shot 所属 camera 的第一个 shot 的首帧路径和描述，用作 reference pair。"""
    camera_tree_path = wd / "camera_tree.json"
    if camera_tree_path.exists():
        cameras = json.loads(camera_tree_path.read_text(encoding="utf-8"))
        for cam in cameras:
            if shot_idx in cam.get("active_shot_idxs", []):
                first_idx = cam["active_shot_idxs"][0]
                ff_path = str(wd / "shots" / str(first_idx) / "first_frame.png")
                desc_path = wd / "shots" / str(first_idx) / "shot_description.json"
                ff_desc = ""
                if desc_path.exists():
                    ff_desc = json.loads(desc_path.read_text(encoding="utf-8")).get("ff_desc", "")
                return (ff_path, ff_desc)
    # fallback：使用自身
    ff_path = str(wd / "shots" / str(shot_idx) / "first_frame.png")
    desc_path = wd / "shots" / str(shot_idx) / "shot_description.json"
    ff_desc = json.loads(desc_path.read_text(encoding="utf-8")).get("ff_desc", "") if desc_path.exists() else ""
    return (ff_path, ff_desc)


def _load_style(wd: Path) -> str:
    meta_path = wd / "metadata.json"
    if meta_path.exists():
        return json.loads(meta_path.read_text(encoding="utf-8")).get("style", "")
    return ""


# ── Frame 启用/禁用状态 ────────────────────────────────────────────────────────

def _abs_to_rel_path(wd: Path, abs_path: str) -> str:
    p = Path(abs_path)
    # Case 1: already an absolute path under wd
    if p.is_absolute():
        try:
            return str(p.relative_to(wd)).replace("\\", "/")
        except ValueError:
            return abs_path
    # Case 2: relative path stored as ".working_dir/{project}/..." (relative to BASE_DIR)
    resolved = BASE_DIR / p
    try:
        return str(resolved.relative_to(wd)).replace("\\", "/")
    except ValueError:
        return abs_path


def _ref_overrides_path(shot_dir: Path, ft_key: str) -> Path:
    return shot_dir / f"{ft_key}_ref_overrides.json"


def _load_ref_overrides(shot_dir: Path, ft_key: str) -> set:
    p = _ref_overrides_path(shot_dir, ft_key)
    if not p.exists():
        return set()
    return set(json.loads(p.read_text(encoding="utf-8")).get("disabled_paths", []))


def _save_ref_overrides(shot_dir: Path, ft_key: str, disabled_paths: set):
    _ref_overrides_path(shot_dir, ft_key).write_text(
        json.dumps({"disabled_paths": sorted(disabled_paths)}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _load_frame_refs(wd: Path, project: str, shot_dir: Path, ft_key: str) -> List[Dict]:
    selector_path = shot_dir / f"{ft_key}_selector_output.json"
    if not selector_path.exists():
        return []
    try:
        selector = json.loads(selector_path.read_text(encoding="utf-8"))
    except Exception:
        return []
    pairs = selector.get("reference_image_path_and_text_pairs", [])
    disabled = _load_ref_overrides(shot_dir, ft_key)
    refs = []
    for abs_path, desc in pairs:
        rel = _abs_to_rel_path(wd, abs_path)
        refs.append({
            "path": abs_path,
            "url": f"/files/{project}/{rel}",
            "description": desc,
            "enabled": abs_path not in disabled,
        })
    return refs


def _frame_state_path(shot_dir: Path) -> Path:
    return shot_dir / "frame_state.json"


def _load_frame_state(shot_dir: Path) -> Dict:
    p = _frame_state_path(shot_dir)
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else {}


def _is_frame_enabled(shot_dir: Path, ft_key: str) -> bool:
    """默认启用（文件存在且未显式禁用）。"""
    return _load_frame_state(shot_dir).get(ft_key, True)


def _set_frame_enabled(shot_dir: Path, ft_key: str, enabled: bool):
    state = _load_frame_state(shot_dir)
    state[ft_key] = enabled
    _frame_state_path(shot_dir).write_text(
        json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8"
    )


# ── 版本管理工具 ──────────────────────────────────────────────────────────────

def _new_vid() -> str:
    return _time.strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:6]


# ── Shot 版本 ──────────────────────────────────────────────────────────────────

def _shot_versions_meta_path(shot_dir: Path) -> Path:
    return shot_dir / "versions.json"


def _load_shot_versions(shot_dir: Path) -> Dict:
    p = _shot_versions_meta_path(shot_dir)
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else {}


def _save_shot_versions(shot_dir: Path, meta: Dict):
    _shot_versions_meta_path(shot_dir).write_text(
        json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _asset_ext(asset: str) -> str:
    return "mp4" if asset == "video" else "png"


def _active_filename(asset: str) -> str:
    return "video.mp4" if asset == "video" else f"{asset}.png"


def _add_shot_version(shot_dir: Path, asset: str, vid: str):
    """将当前活动文件拷贝到 versions 目录，并更新 versions.json（标为已选中）。"""
    ext = _asset_ext(asset)
    active = shot_dir / _active_filename(asset)
    if not active.exists():
        return
    ver_dir = shot_dir / "versions"
    ver_dir.mkdir(exist_ok=True)
    shutil.copy2(str(active), str(ver_dir / f"{asset}_{vid}.{ext}"))

    meta = _load_shot_versions(shot_dir)
    for e in meta.setdefault(asset, []):
        e["selected"] = False
    meta[asset].append({
        "id": vid,
        "created_at": _time.strftime("%Y-%m-%dT%H:%M:%S"),
        "selected": True,
    })
    _save_shot_versions(shot_dir, meta)


def _shot_version_urls(project: str, shot_idx: int, shot_dir: Path) -> Dict:
    meta = _load_shot_versions(shot_dir)
    result = {}
    for asset, entries in meta.items():
        ext = _asset_ext(asset)
        result[asset] = [
            {**e, "url": f"/files/{project}/shots/{shot_idx}/versions/{asset}_{e['id']}.{ext}"}
            for e in entries
        ]
    return result


# ── Portrait 版本 ─────────────────────────────────────────────────────────────

def _portrait_versions_meta_path(char_dir: Path) -> Path:
    return char_dir / "portrait_versions.json"


def _load_portrait_versions(char_dir: Path) -> Dict:
    p = _portrait_versions_meta_path(char_dir)
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else {}


def _save_portrait_versions(char_dir: Path, meta: Dict):
    _portrait_versions_meta_path(char_dir).write_text(
        json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _add_portrait_version(char_dir: Path, view: str, vid: str):
    active = char_dir / f"{view}.png"
    if not active.exists():
        return
    ver_dir = char_dir / "versions"
    ver_dir.mkdir(exist_ok=True)
    shutil.copy2(str(active), str(ver_dir / f"{view}_{vid}.png"))

    meta = _load_portrait_versions(char_dir)
    for e in meta.setdefault(view, []):
        e["selected"] = False
    meta[view].append({
        "id": vid,
        "created_at": _time.strftime("%Y-%m-%dT%H:%M:%S"),
        "selected": True,
    })
    _save_portrait_versions(char_dir, meta)


def _find_char_dir(wd: Path, char_idx: int) -> Optional[Path]:
    portraits_dir = wd / "character_portraits"
    if not portraits_dir.exists():
        return None
    for d in portraits_dir.iterdir():
        if d.is_dir() and d.name.startswith(f"{char_idx}_"):
            return d
    return None


# ── 静态页面 ─────────────────────────────────────────────────────────────────

@app.get("/")
async def index():
    return FileResponse(str(STATIC_DIR / "index.html"))


# ── 项目列表 ─────────────────────────────────────────────────────────────────

@app.get("/api/projects")
async def list_projects():
    if not WORKING_DIR.exists():
        return []
    return sorted(d.name for d in WORKING_DIR.iterdir() if d.is_dir() and not d.name.startswith("."))


# ── 项目数据 ─────────────────────────────────────────────────────────────────

@app.get("/api/projects/{project}/data")
async def get_project_data(project: str):
    wd = _wd(project)

    def load(p: Path):
        return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None

    characters = load(wd / "characters.json") or []
    portraits_registry = load(wd / "character_portraits_registry.json") or {}
    storyboard = load(wd / "storyboard.json") or []
    camera_tree = load(wd / "camera_tree.json") or []
    metadata = load(wd / "metadata.json") or {}

    shots = []
    shots_dir = wd / "shots"
    if shots_dir.exists():
        for shot_dir in sorted(shots_dir.iterdir(), key=lambda d: int(d.name)):
            if not shot_dir.is_dir():
                continue
            idx = int(shot_dir.name)
            desc = load(shot_dir / "shot_description.json")
            scene_refs = []
            refs_dir = shot_dir / "scene_refs"
            if refs_dir.exists():
                for f in sorted(refs_dir.iterdir()):
                    if f.suffix.lower() in SCENE_REF_EXTS:
                        scene_refs.append({"filename": f.name, "path": f"shots/{idx}/scene_refs/{f.name}", "readonly": False})
            for f in sorted(shot_dir.glob("new_camera_*.png")):
                scene_refs.append({"filename": f.name, "path": f"shots/{idx}/{f.name}", "readonly": True})
            has_ff = (shot_dir / "first_frame.png").exists()
            has_lf = (shot_dir / "last_frame.png").exists()
            frame_refs = {}
            for ft_key in ("first_frame", "last_frame"):
                if (shot_dir / f"{ft_key}.png").exists():
                    frame_refs[ft_key] = _load_frame_refs(wd, project, shot_dir, ft_key)
            shots.append({
                "idx": idx,
                "description": desc,
                "has_first_frame": has_ff,
                "has_last_frame": has_lf,
                "has_video": (shot_dir / "video.mp4").exists(),
                "first_frame_enabled": _is_frame_enabled(shot_dir, "first_frame") if has_ff else False,
                "last_frame_enabled": _is_frame_enabled(shot_dir, "last_frame") if has_lf else False,
                "scene_refs": scene_refs,
                "frame_refs": frame_refs,
                "versions": _shot_version_urls(project, idx, shot_dir),
            })

    # 收集各角色画像版本
    portrait_versions: Dict[str, Dict] = {}
    portraits_dir = wd / "character_portraits"
    if portraits_dir.exists():
        for char_dict in characters:
            cidx = char_dict.get("idx")
            cname = char_dict.get("identifier_in_scene", "")
            if cidx is not None:
                char_dir = portraits_dir / f"{cidx}_{cname}"
                pv = _load_portrait_versions(char_dir)
                # 注入 URL
                pv_with_urls: Dict[str, List] = {}
                for view, entries in pv.items():
                    pv_with_urls[view] = [
                        {**e, "url": f"/files/{project}/character_portraits/{cidx}_{cname}/versions/{view}_{e['id']}.png"}
                        for e in entries
                    ]
                portrait_versions[str(cidx)] = pv_with_urls

    return {
        "characters": characters,
        "portraits_registry": portraits_registry,
        "portrait_versions": portrait_versions,
        "storyboard": storyboard,
        "camera_tree": camera_tree,
        "shots": shots,
        "has_final_video": (wd / "final_video.mp4").exists(),
        "metadata": metadata,
    }


# ── 文件代理 ─────────────────────────────────────────────────────────────────

@app.get("/files/{project}/{path:path}")
async def serve_file(project: str, path: str):
    wd = _wd(project)
    file_path = wd / path
    if not file_path.exists():
        raise HTTPException(404, "File not found")
    media_type, _ = mimetypes.guess_type(str(file_path))
    return FileResponse(str(file_path), media_type=media_type or "application/octet-stream")


# ── 替换角色肖像（手动上传） ─────────────────────────────────────────────────

@app.post("/api/projects/{project}/characters/{char_idx}/portraits/{view}")
async def replace_portrait(project: str, char_idx: int, view: str, file: UploadFile):
    if view not in ("front", "side", "back"):
        raise HTTPException(400, "view must be front/side/back")
    wd = _wd(project)
    portraits_dir = wd / "character_portraits"
    target_dir = next((d for d in portraits_dir.iterdir() if d.is_dir() and d.name.startswith(f"{char_idx}_")), None)
    if not target_dir:
        raise HTTPException(404, "Character portrait directory not found")
    _save_upload(file, target_dir / f"{view}.png")
    return {"ok": True}


# ── 替换首帧/末帧（手动上传） ────────────────────────────────────────────────

@app.post("/api/projects/{project}/shots/{shot_idx}/frames/{frame_type}")
async def replace_frame(project: str, shot_idx: int, frame_type: str, file: UploadFile):
    if frame_type not in ("first", "last"):
        raise HTTPException(400, "frame_type must be first/last")
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    if not shot_dir.exists():
        raise HTTPException(404, "Shot directory not found")
    _save_upload(file, shot_dir / f"{frame_type}_frame.png")
    return {"ok": True}


# ── 参考场景：上传 / 删除 ────────────────────────────────────────────────────

@app.post("/api/projects/{project}/shots/{shot_idx}/scene_refs")
async def add_scene_ref(project: str, shot_idx: int, file: UploadFile):
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    if not shot_dir.exists():
        raise HTTPException(404, "Shot directory not found")
    suffix = Path(file.filename or "upload.png").suffix.lower()
    if suffix not in SCENE_REF_EXTS:
        raise HTTPException(400, "Unsupported file type")
    refs_dir = shot_dir / "scene_refs"
    refs_dir.mkdir(exist_ok=True)
    stem = Path(file.filename or "scene").stem
    target = refs_dir / f"{stem}{suffix}"
    counter = 1
    while target.exists():
        target = refs_dir / f"{stem}_{counter}{suffix}"; counter += 1
    _save_upload(file, target)
    return {"ok": True, "filename": target.name, "path": f"shots/{shot_idx}/scene_refs/{target.name}"}


@app.delete("/api/projects/{project}/shots/{shot_idx}/scene_refs/{filename}")
async def delete_scene_ref(project: str, shot_idx: int, filename: str):
    if filename.startswith("new_camera_"):
        raise HTTPException(403, "Auto-generated new_camera files cannot be deleted")
    wd = _wd(project)
    target = wd / "shots" / str(shot_idx) / "scene_refs" / filename
    if not target.exists():
        raise HTTPException(404, "File not found")
    target.unlink()
    return {"ok": True}


# ── 编辑 shot 描述 ────────────────────────────────────────────────────────────

@app.patch("/api/projects/{project}/shots/{shot_idx}/description")
async def update_shot_description(project: str, shot_idx: int, request: Request):
    wd = _wd(project)
    desc_path = wd / "shots" / str(shot_idx) / "shot_description.json"
    if not desc_path.exists():
        raise HTTPException(404, "shot_description.json not found")
    desc = json.loads(desc_path.read_text(encoding="utf-8"))
    body = await request.json()
    editable = ("ff_desc", "lf_desc", "motion_desc", "audio_desc", "visual_desc",
                "ff_vis_char_idxs", "lf_vis_char_idxs")
    for key in editable:
        if key in body:
            desc[key] = body[key]
    desc_path.write_text(json.dumps(desc, ensure_ascii=False, indent=2), encoding="utf-8")
    return desc


@app.patch("/api/projects/{project}/shots/{shot_idx}/frames/{frame_type}/enabled")
async def set_frame_enabled(project: str, shot_idx: int, frame_type: str, request: Request):
    if frame_type not in ("first", "last"):
        raise HTTPException(400, "frame_type must be first/last")
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    ft_key = f"{frame_type}_frame"
    if not (shot_dir / f"{ft_key}.png").exists():
        raise HTTPException(404, f"{ft_key}.png not found")
    body = await request.json()
    enabled = bool(body.get("enabled", True))
    _set_frame_enabled(shot_dir, ft_key, enabled)
    return {"ok": True, "enabled": enabled}


@app.get("/api/projects/{project}/shots/{shot_idx}/frames/{frame_type}/refs")
async def get_frame_refs(project: str, shot_idx: int, frame_type: str):
    if frame_type not in ("first", "last"):
        raise HTTPException(400, "frame_type must be first/last")
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    ft_key = f"{frame_type}_frame"
    return _load_frame_refs(wd, project, shot_dir, ft_key)


@app.patch("/api/projects/{project}/shots/{shot_idx}/frames/{frame_type}/refs")
async def toggle_frame_ref(project: str, shot_idx: int, frame_type: str, request: Request):
    if frame_type not in ("first", "last"):
        raise HTTPException(400, "frame_type must be first/last")
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    ft_key = f"{frame_type}_frame"
    body = await request.json()
    path = body.get("path", "")
    enabled = bool(body.get("enabled", True))
    disabled = _load_ref_overrides(shot_dir, ft_key)
    if enabled:
        disabled.discard(path)
    else:
        disabled.add(path)
    _save_ref_overrides(shot_dir, ft_key, disabled)
    return {"ok": True, "path": path, "enabled": enabled}


@app.delete("/api/projects/{project}/shots/{shot_idx}/frames/{frame_type}")
async def delete_frame(project: str, shot_idx: int, frame_type: str):
    if frame_type not in ("first", "last"):
        raise HTTPException(400, "frame_type must be first/last")
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    ft_key = f"{frame_type}_frame"

    # 删除活动文件
    (shot_dir / f"{ft_key}.png").unlink(missing_ok=True)

    # 清空该 asset 的所有版本文件
    ver_dir = shot_dir / "versions"
    if ver_dir.exists():
        for f in ver_dir.glob(f"{ft_key}_*.png"):
            f.unlink()

    # 更新 versions.json 和 frame_state.json
    meta = _load_shot_versions(shot_dir)
    meta[ft_key] = []
    _save_shot_versions(shot_dir, meta)
    state = _load_frame_state(shot_dir)
    state.pop(ft_key, None)
    _frame_state_path(shot_dir).write_text(
        json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    return {"ok": True}


# ── 脚本文件列表 ─────────────────────────────────────────────────────────────

@app.get("/api/projects/{project}/scripts")
async def list_scripts(project: str):
    wd = _wd(project)
    meta_path = wd / "metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    return meta.get("script_files", [])


@app.post("/api/projects/{project}/scripts")
async def add_script(project: str, request: Request):
    wd = _wd(project)
    body = await request.json()
    path = body.get("path", "").strip()
    if not path or not Path(path).exists():
        raise HTTPException(400, f"File not found: {path}")
    meta_path = wd / "metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    files = meta.setdefault("script_files", [])
    if path not in files:
        files.append(path)
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True, "script_files": files}


@app.delete("/api/projects/{project}/scripts")
async def remove_script(project: str, request: Request):
    wd = _wd(project)
    body = await request.json()
    path = body.get("path", "")
    meta_path = wd / "metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    files = meta.get("script_files", [])
    meta["script_files"] = [f for f in files if f != path]
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True}


@app.get("/api/scripts/parse")
async def parse_script(path: str):
    p = Path(path)
    if not p.exists():
        raise HTTPException(404, f"File not found: {path}")
    suffix = p.suffix.lower()
    if suffix == ".docx":
        return _parse_docx(p)
    elif suffix in (".txt", ".md"):
        return {"type": "text", "filename": p.name, "content": p.read_text(encoding="utf-8")}
    else:
        raise HTTPException(400, f"Unsupported file type: {suffix}")


def _parse_docx(p: Path) -> Dict:
    import docx as _docx
    doc = _docx.Document(str(p))

    # 标题段落
    title = ""
    for para in doc.paragraphs:
        if para.text.strip():
            title = para.text.strip()
            break

    # 解析表格（取第一个表格）
    tables = []
    for tbl in doc.tables:
        rows = []
        for i, row in enumerate(tbl.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append({"is_header": i == 0, "cells": cells})
        tables.append(rows)

    return {
        "type": "docx",
        "filename": p.name,
        "title": title,
        "tables": tables,
    }


# ── 保存 metadata（style 等） ─────────────────────────────────────────────────

@app.patch("/api/projects/{project}/metadata")
async def update_metadata(project: str, request: Request):
    wd = _wd(project)
    meta_path = wd / "metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    body = await request.json()
    meta.update(body)
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return meta


# ── Task 状态查询 ─────────────────────────────────────────────────────────────

@app.get("/api/tasks/{task_id}")
async def get_task(task_id: str):
    if task_id not in _tasks:
        raise HTTPException(404, "Task not found")
    return _tasks[task_id]


# ── AI 重新生成：首帧/末帧 ───────────────────────────────────────────────────

@app.post("/api/projects/{project}/shots/{shot_idx}/regenerate/frame/{frame_type}")
async def regenerate_frame(project: str, shot_idx: int, frame_type: str):
    if frame_type not in ("first", "last"):
        raise HTTPException(400, "frame_type must be first/last")
    wd = _wd(project)
    desc_path = wd / "shots" / str(shot_idx) / "shot_description.json"
    if not desc_path.exists():
        raise HTTPException(404, "shot_description.json not found")

    pipeline = get_pipeline(project)
    tid = _new_task()

    asyncio.create_task(_run_regenerate_frame(tid, pipeline, wd, shot_idx, frame_type))
    return {"task_id": tid}


async def _run_regenerate_frame(tid: str, pipeline, wd: Path, shot_idx: int, frame_type: str):
    try:
        from interfaces import CharacterInScene, ShotDescription

        ft_key = f"{frame_type}_frame"  # "first_frame" or "last_frame"

        _task_progress(tid, "归档旧版本…")
        shot_dir = wd / "shots" / str(shot_idx)
        frame_path = shot_dir / f"{ft_key}.png"
        selector_path = shot_dir / f"{ft_key}_selector_output.json"

        # 先把当前文件归档到 versions，这样生成期间 UI 仍能展示旧图
        old_vid = None
        if frame_path.exists():
            old_vid = _new_vid()
            _add_shot_version(shot_dir, ft_key, old_vid)
            _tasks[tid]["preview_url"] = (
                f"/files/{wd.name}/shots/{shot_idx}/versions/{ft_key}_{old_vid}.png"
            )

        # 删除活动文件和 selector（强制 pipeline 重新生成）
        frame_path.unlink(missing_ok=True)
        selector_path.unlink(missing_ok=True)
        _task_progress(tid, "加载上下文…")

        # 加载上下文
        desc = ShotDescription.model_validate(
            json.loads((wd / "shots" / str(shot_idx) / "shot_description.json").read_text(encoding="utf-8"))
        )
        characters = [
            CharacterInScene.model_validate(c)
            for c in json.loads((wd / "characters.json").read_text(encoding="utf-8"))
        ]
        registry = json.loads((wd / "character_portraits_registry.json").read_text(encoding="utf-8"))

        vis_idxs = desc.ff_vis_char_idxs if frame_type == "first" else desc.lf_vis_char_idxs
        visible_chars = [c for c in characters if c.idx in vis_idxs]
        frame_desc = desc.ff_desc if frame_type == "first" else desc.lf_desc
        first_shot_ff_pair = _get_first_shot_ff_pair(wd, shot_idx)

        # 若 reference 指向被删掉的文件（自身是 camera 第一帧的情况），改用归档版本；
        # 若归档也不存在（首次生成），则不传 reference，避免 FileNotFoundError。
        ff_ref_path, ff_ref_text = first_shot_ff_pair
        if not Path(ff_ref_path).exists():
            archived = shot_dir / "versions" / f"{ft_key}_{old_vid}.png" if old_vid is not None else None
            if archived is not None and archived.exists():
                first_shot_ff_pair = (str(archived), ff_ref_text)
            else:
                first_shot_ff_pair = None

        _task_progress(tid, "AI 筛选参考图 + 生成 Prompt…")
        # 预初始化 frame_events
        _setup_frame_events(pipeline, wd, shot_idx, ft_key)

        _task_progress(tid, "图像生成 API 调用中…")
        variation_hint = (
            "IMPORTANT: Generate a noticeably different visual composition compared to any previous version. "
            "Explore a different camera framing, lighting direction, character pose, or spatial arrangement "
            "while keeping the scene description accurate and character appearances consistent with the reference images."
        )
        excluded_ref_paths = list(_load_ref_overrides(shot_dir, ft_key))
        await pipeline.generate_frame_for_single_shot(
            shot_idx=shot_idx,
            frame_type=ft_key,
            first_shot_ff_path_and_text_pair=first_shot_ff_pair,
            frame_desc=frame_desc,
            visible_characters=visible_chars,
            character_portraits_registry=registry,
            prompt_suffix=variation_hint,
            excluded_ref_paths=excluded_ref_paths,
        )
        # 归档版本，并确保启用状态
        _task_progress(tid, "归档版本…")
        vid = _new_vid()
        _add_shot_version(wd / "shots" / str(shot_idx), ft_key, vid)
        _set_frame_enabled(wd / "shots" / str(shot_idx), ft_key, True)
        _tasks[tid]["new_vid"] = vid
        _tasks[tid]["enabled"] = True
        _task_done(tid)
    except Exception as e:
        logger.exception(f"Task {tid} failed")
        _task_error(tid, str(e))


# ── AI 重新生成：视频 ────────────────────────────────────────────────────────

@app.post("/api/projects/{project}/shots/{shot_idx}/regenerate/video")
async def regenerate_video(project: str, shot_idx: int):
    wd = _wd(project)
    desc_path = wd / "shots" / str(shot_idx) / "shot_description.json"
    if not desc_path.exists():
        raise HTTPException(404, "shot_description.json not found")

    pipeline = get_pipeline(project)
    tid = _new_task()
    asyncio.create_task(_run_regenerate_video(tid, pipeline, wd, shot_idx))
    return {"task_id": tid}


async def _run_regenerate_video(tid: str, pipeline, wd: Path, shot_idx: int):
    try:
        from interfaces import ShotDescription
        from tools.protocols import VideoRAIFilteredError

        _task_progress(tid, "归档旧版本…")
        shot_dir = wd / "shots" / str(shot_idx)
        video_path = shot_dir / "video.mp4"

        if video_path.exists():
            old_vid = _new_vid()
            _add_shot_version(shot_dir, "video", old_vid)
            _tasks[tid]["preview_url"] = (
                f"/files/{wd.name}/shots/{shot_idx}/versions/video_{old_vid}.mp4"
            )

        video_path.unlink(missing_ok=True)
        _task_progress(tid, "加载镜头描述…")

        desc = ShotDescription.model_validate(
            json.loads((wd / "shots" / str(shot_idx) / "shot_description.json").read_text(encoding="utf-8"))
        )

        # 只把启用（enabled）的帧传入，空列表 → 文生视频（t2v）
        frame_paths = []
        ff = shot_dir / "first_frame.png"
        lf = shot_dir / "last_frame.png"
        if ff.exists() and _is_frame_enabled(shot_dir, "first_frame"):
            frame_paths.append(str(ff))
        if lf.exists() and _is_frame_enabled(shot_dir, "last_frame") and desc.variation_type in ("medium", "large"):
            frame_paths.append(str(lf))

        mode = {0: "文生视频", 1: "首帧→视频", 2: "首尾帧→视频"}[len(frame_paths)]
        _task_progress(tid, f"视频生成 API 调用中（{mode}）…")
        logger.info(f"Shot {shot_idx} video regeneration: {mode}")

        try:
            video_output = await pipeline.video_generator.generate_single_video(
                prompt=desc.motion_desc + "\n" + desc.audio_desc,
                reference_image_paths=frame_paths,
            )
        except VideoRAIFilteredError:
            logger.warning(f"Shot {shot_idx}: audio_desc filtered by RAI, retrying with motion_desc only")
            _task_progress(tid, "RAI 过滤，仅用运动描述重试…")
            video_output = await pipeline.video_generator.generate_single_video(
                prompt=desc.motion_desc,
                reference_image_paths=frame_paths,
            )

        video_output.save(str(video_path))
        # 归档版本
        vid = _new_vid()
        _add_shot_version(wd / "shots" / str(shot_idx), "video", vid)
        _tasks[tid]["new_vid"] = vid
        _task_done(tid)
    except Exception as e:
        logger.exception(f"Task {tid} failed")
        _task_error(tid, str(e))


# ── AI 重新生成：角色肖像 ────────────────────────────────────────────────────

@app.post("/api/projects/{project}/characters/{char_idx}/portraits/{view}/regenerate")
async def regenerate_portrait(project: str, char_idx: int, view: str):
    if view not in ("front", "side", "back"):
        raise HTTPException(400, "view must be front/side/back")
    wd = _wd(project)
    chars_path = wd / "characters.json"
    if not chars_path.exists():
        raise HTTPException(404, "characters.json not found")

    pipeline = get_pipeline(project)
    tid = _new_task()
    asyncio.create_task(_run_regenerate_portrait(tid, pipeline, wd, char_idx, view))
    return {"task_id": tid}


async def _run_regenerate_portrait(tid: str, pipeline, wd: Path, char_idx: int, view: str):
    try:
        from interfaces import CharacterInScene
        from pipelines.script2video_pipeline import Script2VideoPipeline

        _task_progress(tid, "加载角色信息…")
        characters = [
            CharacterInScene.model_validate(c)
            for c in json.loads((wd / "characters.json").read_text(encoding="utf-8"))
        ]
        char = next((c for c in characters if c.idx == char_idx), None)
        if not char:
            raise ValueError(f"Character idx={char_idx} not found")

        char_dir = wd / "character_portraits" / f"{char_idx}_{char.identifier_in_scene}"
        front_path = char_dir / "front.png"
        style = _load_style(wd)

        # 预初始化 character_portrait_events
        Script2VideoPipeline.character_portrait_events = {}
        ev = asyncio.Event()
        Script2VideoPipeline.character_portrait_events[char_idx] = ev

        # 归档旧版本，生成期间 UI 继续展示旧图
        portrait_path = char_dir / f"{view}.png"
        if portrait_path.exists():
            old_vid = _new_vid()
            _add_portrait_version(char_dir, view, old_vid)
            _tasks[tid]["preview_url"] = (
                f"/files/{wd.name}/character_portraits/{char_dir.name}/versions/{view}_{old_vid}.png"
            )

        if view == "front":
            _task_progress(tid, "图像生成 API 调用中（正面肖像）…")
            front_path.unlink(missing_ok=True)
            output = await pipeline.character_portraits_generator.generate_front_portrait(char, style)
            output.save(str(front_path))

        elif view == "side":
            if not front_path.exists():
                raise ValueError("front portrait must exist to regenerate side portrait")
            _task_progress(tid, "图像生成 API 调用中（侧面肖像）…")
            side_path = char_dir / "side.png"
            side_path.unlink(missing_ok=True)
            output = await pipeline.character_portraits_generator.generate_side_portrait(char, str(front_path))
            output.save(str(side_path))

        elif view == "back":
            if not front_path.exists():
                raise ValueError("front portrait must exist to regenerate back portrait")
            _task_progress(tid, "图像生成 API 调用中（背面肖像）…")
            back_path = char_dir / "back.png"
            back_path.unlink(missing_ok=True)
            output = await pipeline.character_portraits_generator.generate_back_portrait(char, str(front_path))
            output.save(str(back_path))

        ev.set()
        # 归档版本
        vid = _new_vid()
        _add_portrait_version(char_dir, view, vid)
        _tasks[tid]["new_vid"] = vid
        _task_done(tid)
    except Exception as e:
        logger.exception(f"Task {tid} failed")
        _task_error(tid, str(e))


# ── Shot 版本管理 ─────────────────────────────────────────────────────────────

@app.get("/api/projects/{project}/shots/{shot_idx}/versions")
async def get_shot_versions(project: str, shot_idx: int):
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    return _shot_version_urls(project, shot_idx, shot_dir)


@app.post("/api/projects/{project}/shots/{shot_idx}/versions/{asset}/{vid}/select")
async def select_shot_version(project: str, shot_idx: int, asset: str, vid: str):
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    ext = _asset_ext(asset)
    ver_file = shot_dir / "versions" / f"{asset}_{vid}.{ext}"
    if not ver_file.exists():
        raise HTTPException(404, "Version file not found")
    shutil.copy2(str(ver_file), str(shot_dir / _active_filename(asset)))
    meta = _load_shot_versions(shot_dir)
    for e in meta.get(asset, []):
        e["selected"] = e["id"] == vid
    _save_shot_versions(shot_dir, meta)
    return {"ok": True}


@app.delete("/api/projects/{project}/shots/{shot_idx}/versions/{asset}/{vid}")
async def delete_shot_version(project: str, shot_idx: int, asset: str, vid: str):
    wd = _wd(project)
    shot_dir = wd / "shots" / str(shot_idx)
    ext = _asset_ext(asset)
    ver_file = shot_dir / "versions" / f"{asset}_{vid}.{ext}"
    ver_file.unlink(missing_ok=True)

    meta = _load_shot_versions(shot_dir)
    entries = meta.get(asset, [])
    was_selected = any(e["id"] == vid and e["selected"] for e in entries)
    meta[asset] = [e for e in entries if e["id"] != vid]

    if was_selected:
        if meta[asset]:
            # 自动选中最近一个
            meta[asset][-1]["selected"] = True
            last_id = meta[asset][-1]["id"]
            shutil.copy2(
                str(shot_dir / "versions" / f"{asset}_{last_id}.{ext}"),
                str(shot_dir / _active_filename(asset)),
            )
        else:
            (shot_dir / _active_filename(asset)).unlink(missing_ok=True)

    _save_shot_versions(shot_dir, meta)
    return {"ok": True, "was_selected": was_selected, "versions": _shot_version_urls(project, shot_idx, shot_dir)}


# ── Portrait 版本管理 ─────────────────────────────────────────────────────────

@app.get("/api/projects/{project}/characters/{char_idx}/portraits/versions")
async def get_portrait_versions_all(project: str, char_idx: int):
    wd = _wd(project)
    char_dir = _find_char_dir(wd, char_idx)
    if not char_dir:
        raise HTTPException(404, "Character portrait directory not found")
    meta = _load_portrait_versions(char_dir)
    result = {}
    for view, entries in meta.items():
        result[view] = [
            {**e, "url": f"/files/{project}/character_portraits/{char_dir.name}/versions/{view}_{e['id']}.png"}
            for e in entries
        ]
    return result


@app.post("/api/projects/{project}/characters/{char_idx}/portraits/{view}/versions/{vid}/select")
async def select_portrait_version(project: str, char_idx: int, view: str, vid: str):
    wd = _wd(project)
    char_dir = _find_char_dir(wd, char_idx)
    if not char_dir:
        raise HTTPException(404, "Character portrait directory not found")
    ver_file = char_dir / "versions" / f"{view}_{vid}.png"
    if not ver_file.exists():
        raise HTTPException(404, "Version file not found")
    shutil.copy2(str(ver_file), str(char_dir / f"{view}.png"))
    meta = _load_portrait_versions(char_dir)
    for e in meta.get(view, []):
        e["selected"] = e["id"] == vid
    _save_portrait_versions(char_dir, meta)
    return {"ok": True}


@app.delete("/api/projects/{project}/characters/{char_idx}/portraits/{view}/versions/{vid}")
async def delete_portrait_version(project: str, char_idx: int, view: str, vid: str):
    wd = _wd(project)
    char_dir = _find_char_dir(wd, char_idx)
    if not char_dir:
        raise HTTPException(404, "Character portrait directory not found")
    ver_file = char_dir / "versions" / f"{view}_{vid}.png"
    ver_file.unlink(missing_ok=True)

    meta = _load_portrait_versions(char_dir)
    entries = meta.get(view, [])
    was_selected = any(e["id"] == vid and e["selected"] for e in entries)
    meta[view] = [e for e in entries if e["id"] != vid]

    if was_selected:
        if meta[view]:
            meta[view][-1]["selected"] = True
            last_id = meta[view][-1]["id"]
            shutil.copy2(
                str(char_dir / "versions" / f"{view}_{last_id}.png"),
                str(char_dir / f"{view}.png"),
            )
        else:
            (char_dir / f"{view}.png").unlink(missing_ok=True)

    _save_portrait_versions(char_dir, meta)
    result = {}
    updated = _load_portrait_versions(char_dir)
    for v, es in updated.items():
        result[v] = [
            {**e, "url": f"/files/{project}/character_portraits/{char_dir.name}/versions/{v}_{e['id']}.png"}
            for e in es
        ]
    return {"ok": True, "was_selected": was_selected, "versions": result}


# ── 阶段状态 & 手动触发 ────────────────────────────────────────────────────────

def _get_script(wd: Path) -> str:
    """从 metadata.script_files 读取第一个脚本文本；找不到则返回空字符串。"""
    meta_path = wd / "metadata.json"
    if not meta_path.exists():
        return ""
    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    files = meta.get("script_files", [])
    if not files:
        return ""
    p = Path(files[0])
    if not p.exists():
        return ""
    suffix = p.suffix.lower()
    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8")
    if suffix == ".docx":
        try:
            import docx as _docx
            doc = _docx.Document(str(p))
            return "\n".join(para.text for para in doc.paragraphs)
        except Exception:
            return ""
    return ""


def _stages_status(wd: Path) -> Dict:
    chars_path = wd / "characters.json"
    registry_path = wd / "character_portraits_registry.json"
    storyboard_path = wd / "storyboard.json"
    camera_tree_path = wd / "camera_tree.json"
    final_video_path = wd / "final_video.mp4"

    chars: List = []
    if chars_path.exists():
        try:
            chars = json.loads(chars_path.read_text(encoding="utf-8"))
        except Exception:
            pass

    portrait_count = 0
    if registry_path.exists():
        try:
            reg = json.loads(registry_path.read_text(encoding="utf-8"))
            portrait_count = sum(len(v) for v in reg.values())
        except Exception:
            pass

    storyboard: List = []
    if storyboard_path.exists():
        try:
            storyboard = json.loads(storyboard_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    storyboard_count = len(storyboard)

    shots_dir = wd / "shots"
    shot_desc_count = 0
    shots_with_ff = 0
    shots_with_video = 0
    if shots_dir.exists():
        for d in shots_dir.iterdir():
            if d.is_dir():
                if (d / "shot_description.json").exists():
                    shot_desc_count += 1
                if (d / "first_frame.png").exists():
                    shots_with_ff += 1
                if (d / "video.mp4").exists():
                    shots_with_video += 1

    return {
        "extract_characters": {
            "done": chars_path.exists(),
            "count": len(chars),
        },
        "generate_portraits": {
            "done": registry_path.exists() and portrait_count > 0,
            "count": portrait_count,
        },
        "design_storyboard": {
            "done": storyboard_path.exists() and storyboard_count > 0,
            "count": storyboard_count,
        },
        "decompose_descriptions": {
            "done": storyboard_count > 0 and shot_desc_count == storyboard_count,
            "shots_done": shot_desc_count,
            "shots_total": storyboard_count,
        },
        "construct_camera_tree": {
            "done": camera_tree_path.exists(),
        },
        "generate_frames": {
            "done": storyboard_count > 0 and shots_with_ff == storyboard_count,
            "shots_done": shots_with_ff,
            "shots_total": storyboard_count,
        },
        "generate_videos": {
            "done": storyboard_count > 0 and shots_with_video == storyboard_count,
            "shots_done": shots_with_video,
            "shots_total": storyboard_count,
        },
        "concatenate": {
            "done": final_video_path.exists(),
        },
    }


def _clear_stage_cache(wd: Path, stage: str):
    """删除指定阶段的缓存文件，以便强制重新运行该阶段。"""
    shots_dir = wd / "shots"

    if stage == "extract_characters":
        (wd / "characters.json").unlink(missing_ok=True)

    elif stage == "generate_portraits":
        (wd / "character_portraits_registry.json").unlink(missing_ok=True)
        portraits_dir = wd / "character_portraits"
        if portraits_dir.exists():
            shutil.rmtree(str(portraits_dir))

    elif stage == "design_storyboard":
        (wd / "storyboard.json").unlink(missing_ok=True)

    elif stage == "decompose_descriptions":
        if shots_dir.exists():
            for d in shots_dir.iterdir():
                if d.is_dir():
                    (d / "shot_description.json").unlink(missing_ok=True)

    elif stage == "construct_camera_tree":
        (wd / "camera_tree.json").unlink(missing_ok=True)

    elif stage == "generate_frames":
        if shots_dir.exists():
            for d in shots_dir.iterdir():
                if d.is_dir():
                    for fname in (
                        "first_frame.png", "last_frame.png",
                        "first_frame_selector_output.json", "last_frame_selector_output.json",
                        "first_frame_ref_overrides.json", "last_frame_ref_overrides.json",
                    ):
                        (d / fname).unlink(missing_ok=True)
                    for f in d.glob("transition_video_from_shot_*.mp4"):
                        f.unlink(missing_ok=True)
                    for f in d.glob("new_camera_*.png"):
                        f.unlink(missing_ok=True)

    elif stage == "generate_videos":
        if shots_dir.exists():
            for d in shots_dir.iterdir():
                if d.is_dir():
                    (d / "video.mp4").unlink(missing_ok=True)

    elif stage == "concatenate":
        (wd / "final_video.mp4").unlink(missing_ok=True)


_STAGE_PREREQUISITES: Dict[str, List[str]] = {
    "extract_characters":     [],
    "generate_portraits":     ["characters.json"],
    "design_storyboard":      ["characters.json"],
    "decompose_descriptions":  ["storyboard.json"],
    "construct_camera_tree":  [],   # 动态检查 shot_description 文件
    "generate_frames":        ["camera_tree.json", "character_portraits_registry.json"],
    "generate_videos":        [],   # 动态检查 first_frame 文件
    "concatenate":            [],   # 动态检查 video 文件
}


@app.get("/api/projects/{project}/stages/status")
async def get_stages_status(project: str):
    wd = _wd(project)
    return _stages_status(wd)


@app.post("/api/projects/{project}/run/{stage}")
async def run_stage(project: str, stage: str, request: Request):
    valid_stages = set(_STAGE_PREREQUISITES.keys())
    if stage not in valid_stages:
        raise HTTPException(400, f"Unknown stage '{stage}'. Valid: {sorted(valid_stages)}")

    wd = _wd(project)
    body: Dict = {}
    try:
        body = await request.json()
    except Exception:
        pass
    force = bool(body.get("force", False))

    # 静态前置检查
    for prereq in _STAGE_PREREQUISITES[stage]:
        if not (wd / prereq).exists():
            raise HTTPException(400, f"前置条件未满足：{prereq} 不存在，请先运行依赖阶段")

    # 动态前置检查
    storyboard_loaded: Optional[List] = None

    def _load_storyboard() -> List:
        nonlocal storyboard_loaded
        if storyboard_loaded is None:
            p = wd / "storyboard.json"
            if not p.exists():
                raise HTTPException(400, "前置条件未满足：storyboard.json 不存在，请先运行 design_storyboard")
            storyboard_loaded = json.loads(p.read_text(encoding="utf-8"))
        return storyboard_loaded

    if stage == "construct_camera_tree":
        storyboard = _load_storyboard()
        shots_dir = wd / "shots"
        for shot in storyboard:
            idx = shot["idx"]
            if not (shots_dir / str(idx) / "shot_description.json").exists():
                raise HTTPException(
                    400,
                    f"前置条件未满足：shot {idx} 的 shot_description.json 不存在，请先运行 decompose_descriptions",
                )

    elif stage == "generate_videos":
        storyboard = _load_storyboard()
        shots_dir = wd / "shots"
        for shot in storyboard:
            idx = shot["idx"]
            if not (shots_dir / str(idx) / "first_frame.png").exists():
                raise HTTPException(
                    400,
                    f"前置条件未满足：shot {idx} 的 first_frame.png 不存在，请先运行 generate_frames",
                )

    elif stage == "concatenate":
        storyboard = _load_storyboard()
        shots_dir = wd / "shots"
        for shot in storyboard:
            idx = shot["idx"]
            if not (shots_dir / str(idx) / "video.mp4").exists():
                raise HTTPException(
                    400,
                    f"前置条件未满足：shot {idx} 的 video.mp4 不存在，请先运行 generate_videos",
                )

    if force:
        _clear_stage_cache(wd, stage)

    pipeline = get_pipeline(project)
    tid = _new_task()
    asyncio.create_task(_run_stage_task(tid, pipeline, wd, stage))
    return {"task_id": tid}


async def _run_stage_task(tid: str, pipeline, wd: Path, stage: str):
    try:
        if stage == "extract_characters":
            await _stage_extract_characters(tid, pipeline, wd)
        elif stage == "generate_portraits":
            await _stage_generate_portraits(tid, pipeline, wd)
        elif stage == "design_storyboard":
            await _stage_design_storyboard(tid, pipeline, wd)
        elif stage == "decompose_descriptions":
            await _stage_decompose_descriptions(tid, pipeline, wd)
        elif stage == "construct_camera_tree":
            await _stage_construct_camera_tree(tid, pipeline, wd)
        elif stage == "generate_frames":
            await _stage_generate_frames(tid, pipeline, wd)
        elif stage == "generate_videos":
            await _stage_generate_videos(tid, pipeline, wd)
        elif stage == "concatenate":
            await _stage_concatenate(tid, wd)
        _task_done(tid)
    except Exception as e:
        logger.exception(f"Stage task {tid} ({stage}) failed")
        _task_error(tid, str(e))


async def _stage_extract_characters(tid: str, pipeline, wd: Path):
    _task_progress(tid, "读取脚本…")
    script = _get_script(wd)
    if not script:
        raise ValueError("找不到脚本文件，请先在项目中添加脚本（scripts）")
    _task_progress(tid, "LLM 提取角色中…")
    characters = await pipeline.extract_characters(script=script)
    _task_progress(tid, f"完成，提取到 {len(characters)} 个角色")


async def _stage_generate_portraits(tid: str, pipeline, wd: Path):
    from interfaces import CharacterInScene
    _task_progress(tid, "加载角色信息…")
    characters = [
        CharacterInScene.model_validate(c)
        for c in json.loads((wd / "characters.json").read_text(encoding="utf-8"))
    ]
    style = _load_style(wd)
    _task_progress(tid, f"生成 {len(characters)} 个角色的肖像（front/side/back）…")
    await pipeline.generate_character_portraits(
        characters=characters,
        character_portraits_registry=None,
        style=style,
    )
    _task_progress(tid, "角色肖像生成完成")


async def _stage_design_storyboard(tid: str, pipeline, wd: Path):
    from interfaces import CharacterInScene
    _task_progress(tid, "读取脚本和角色…")
    script = _get_script(wd)
    if not script:
        raise ValueError("找不到脚本文件，请先在项目中添加脚本（scripts）")
    characters = [
        CharacterInScene.model_validate(c)
        for c in json.loads((wd / "characters.json").read_text(encoding="utf-8"))
    ]
    meta: Dict = {}
    meta_path = wd / "metadata.json"
    if meta_path.exists():
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
    user_requirement = meta.get("user_requirement", "")
    _task_progress(tid, "LLM 设计分镜中…")
    storyboard = await pipeline.design_storyboard(
        script=script,
        characters=characters,
        user_requirement=user_requirement,
    )
    _task_progress(tid, f"分镜设计完成，共 {len(storyboard)} 个镜头")


async def _stage_decompose_descriptions(tid: str, pipeline, wd: Path):
    from interfaces import CharacterInScene, ShotBriefDescription
    _task_progress(tid, "加载分镜和角色…")
    characters = [
        CharacterInScene.model_validate(c)
        for c in json.loads((wd / "characters.json").read_text(encoding="utf-8"))
    ]
    storyboard = [
        ShotBriefDescription.model_validate(s)
        for s in json.loads((wd / "storyboard.json").read_text(encoding="utf-8"))
    ]
    _task_progress(tid, f"并发拆解 {len(storyboard)} 个镜头的视觉描述…")
    shot_descriptions = await pipeline.decompose_visual_descriptions(
        shot_brief_descriptions=storyboard,
        characters=characters,
    )
    _task_progress(tid, f"视觉描述拆解完成，共 {len(shot_descriptions)} 个镜头")


async def _stage_construct_camera_tree(tid: str, pipeline, wd: Path):
    from interfaces import ShotDescription
    _task_progress(tid, "加载镜头描述…")
    shots_dir = wd / "shots"
    shot_descriptions = [
        ShotDescription.model_validate(
            json.loads((d / "shot_description.json").read_text(encoding="utf-8"))
        )
        for d in sorted(shots_dir.iterdir(), key=lambda x: int(x.name))
        if d.is_dir() and (d / "shot_description.json").exists()
    ]
    _task_progress(tid, f"LLM 构建相机树（{len(shot_descriptions)} 个镜头）…")
    camera_tree = await pipeline.construct_camera_tree(shot_descriptions=shot_descriptions)
    _task_progress(tid, f"相机树构建完成，共 {len(camera_tree)} 个相机")


async def _stage_generate_frames(tid: str, pipeline, wd: Path):
    from interfaces import CharacterInScene, ShotDescription, Camera
    from pipelines.script2video_pipeline import Script2VideoPipeline

    _task_progress(tid, "加载上下文…")
    characters = [
        CharacterInScene.model_validate(c)
        for c in json.loads((wd / "characters.json").read_text(encoding="utf-8"))
    ]
    registry = json.loads((wd / "character_portraits_registry.json").read_text(encoding="utf-8"))
    camera_tree = [
        Camera.model_validate(c)
        for c in json.loads((wd / "camera_tree.json").read_text(encoding="utf-8"))
    ]
    shots_dir = wd / "shots"
    shot_descriptions = [
        ShotDescription.model_validate(
            json.loads((d / "shot_description.json").read_text(encoding="utf-8"))
        )
        for d in sorted(shots_dir.iterdir(), key=lambda x: int(x.name))
        if d.is_dir() and (d / "shot_description.json").exists()
    ]

    # 初始化 frame_events（全部 unset，让 pipeline 正常完成后 set）
    Script2VideoPipeline.frame_events = {
        sd.idx: {"first_frame": asyncio.Event(), "last_frame": asyncio.Event()}
        for sd in shot_descriptions
    }

    priority_shot_idxs = [cam.parent_cam_idx for cam in camera_tree if cam.parent_cam_idx is not None]
    _task_progress(tid, f"并发生成 {len(camera_tree)} 个相机的帧图…")
    await asyncio.gather(*[
        pipeline.generate_frames_for_single_camera(
            camera=cam,
            shot_descriptions=shot_descriptions,
            characters=characters,
            character_portraits_registry=registry,
            priority_shot_idxs=priority_shot_idxs,
        )
        for cam in camera_tree
    ])
    _task_progress(tid, "帧图生成完成")


async def _stage_generate_videos(tid: str, pipeline, wd: Path):
    from interfaces import ShotDescription

    _task_progress(tid, "加载镜头描述…")
    shots_dir = wd / "shots"
    shot_descriptions = [
        ShotDescription.model_validate(
            json.loads((d / "shot_description.json").read_text(encoding="utf-8"))
        )
        for d in sorted(shots_dir.iterdir(), key=lambda x: int(x.name))
        if d.is_dir() and (d / "shot_description.json").exists()
    ]

    _preset_all_frame_events(pipeline, wd)

    _task_progress(tid, f"并发生成 {len(shot_descriptions)} 个镜头的视频…")
    await asyncio.gather(*[
        pipeline.generate_video_for_single_shot(shot_description=sd)
        for sd in shot_descriptions
    ])
    _task_progress(tid, "视频生成完成")


async def _stage_concatenate(tid: str, wd: Path):
    from moviepy import VideoFileClip, concatenate_videoclips

    _task_progress(tid, "加载分镜顺序…")
    storyboard = json.loads((wd / "storyboard.json").read_text(encoding="utf-8"))
    shots_dir = wd / "shots"

    _task_progress(tid, f"拼接 {len(storyboard)} 个镜头…")
    clips = [
        VideoFileClip(str(shots_dir / str(s["idx"]) / "video.mp4"))
        for s in storyboard
    ]
    final = concatenate_videoclips(clips)
    final_path = wd / "final_video.mp4"
    final.write_videofile(str(final_path), codec="libx264", preset="medium")
    _task_progress(tid, f"最终视频已保存：{final_path.name}")


# ── 内部工具 ─────────────────────────────────────────────────────────────────

def _wd(project: str) -> Path:
    wd = WORKING_DIR / project
    if not wd.exists():
        raise HTTPException(404, f"Project '{project}' not found")
    return wd


def _save_upload(file: UploadFile, target_path: Path):
    target_path.parent.mkdir(parents=True, exist_ok=True)
    with target_path.open("wb") as f:
        shutil.copyfileobj(file.file, f)
