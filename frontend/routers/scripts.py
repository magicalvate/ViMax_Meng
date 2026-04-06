import json
from pathlib import Path

from fastapi import APIRouter, HTTPException, Request, UploadFile, File

from frontend.core import wd

router = APIRouter()


@router.get("/api/projects/{project}/scripts")
async def list_scripts(project: str):
    p = wd(project)
    meta_path = p / "metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    return meta.get("script_files", [])


@router.post("/api/projects/{project}/scripts")
async def add_script(project: str, request: Request):
    p = wd(project)
    body = await request.json()
    path = body.get("path", "").strip()
    if not path or not Path(path).exists():
        raise HTTPException(400, f"File not found: {path}")
    meta_path = p / "metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    files = meta.setdefault("script_files", [])
    if path not in files:
        files.append(path)
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True, "script_files": files}


@router.post("/api/projects/{project}/scripts/upload")
async def upload_script(project: str, file: UploadFile = File(...)):
    p = wd(project)
    scripts_dir = p / "scripts"
    scripts_dir.mkdir(exist_ok=True)
    dest = scripts_dir / file.filename
    content = await file.read()
    dest.write_bytes(content)
    # add to metadata
    script_path = str(dest)
    meta_path = p / "metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    files = meta.setdefault("script_files", [])
    if script_path not in files:
        files.append(script_path)
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True, "script_files": files, "path": script_path}


@router.delete("/api/projects/{project}/scripts")
async def remove_script(project: str, request: Request):
    p = wd(project)
    body = await request.json()
    path = body.get("path", "")
    meta_path = p / "metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    meta["script_files"] = [f for f in meta.get("script_files", []) if f != path]
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True}


@router.get("/api/scripts/parse")
async def parse_script(path: str):
    p = Path(path)
    if not p.exists():
        raise HTTPException(404, f"File not found: {path}")
    suffix = p.suffix.lower()
    if suffix == ".docx":
        return _parse_docx(p)
    elif suffix in (".txt", ".md"):
        return {"type": "text", "filename": p.name, "content": p.read_text(encoding="utf-8")}
    else:
        raise HTTPException(400, f"Unsupported file type: {suffix}")


def _parse_docx(p: Path) -> dict:
    import docx as _docx
    doc = _docx.Document(str(p))
    title = ""
    for para in doc.paragraphs:
        if para.text.strip():
            title = para.text.strip()
            break
    tables = []
    for tbl in doc.tables:
        rows = []
        for i, row in enumerate(tbl.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append({"is_header": i == 0, "cells": cells})
        tables.append(rows)
    return {"type": "docx", "filename": p.name, "title": title, "tables": tables}
